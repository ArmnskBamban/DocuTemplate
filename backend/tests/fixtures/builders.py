"""Programmatic DOCX fixture builders (spec Section 63).

Fixtures are generated at test time instead of committed as binary blobs. This
keeps tests readable and allows us to assert formatting/layout properties.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

OLD_BACKGROUND = "OLD BACKGROUND CONTENT UNIQUE_SECRET_REPORT_SENTENCE_123"
OLD_OBJECTIVES = "OLD OBJECTIVES"
OLD_THEORY = "OLD THEORY"
OLD_RESULT = "OLD RESULT CONTENT"
OLD_CONCLUSION = "OLD CONCLUSION"


def build_acceptance_docx(path: Path) -> Path:
    """Fixture matching the acceptance scenario (spec Section 74)."""
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(3)
    sec.bottom_margin = Cm(3)
    sec.left_margin = Cm(4)
    sec.right_margin = Cm(3)

    doc.add_paragraph("UNIVERSITAS ANDALAS").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("LAPORAN PRAKTIKUM").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Nama: John Doe")
    doc.add_paragraph("NIM: 24123456")

    doc.add_heading("BAB I", level=1)
    doc.add_heading("PENDAHULUAN", level=1)
    doc.add_heading("1.1 Latar Belakang", level=2)
    doc.add_paragraph(OLD_BACKGROUND)
    doc.add_heading("1.2 Tujuan", level=2)
    doc.add_paragraph(OLD_OBJECTIVES)

    doc.add_heading("BAB II", level=1)
    doc.add_heading("LANDASAN TEORI", level=1)
    doc.add_paragraph(OLD_THEORY)

    doc.add_heading("BAB III", level=1)
    doc.add_heading("HASIL", level=1)
    doc.add_paragraph(OLD_RESULT)

    doc.add_heading("BAB IV", level=1)
    doc.add_heading("KESIMPULAN", level=1)
    doc.add_paragraph(OLD_CONCLUSION)

    doc.save(path)
    return path


def build_custom_heading_docx(path: Path) -> Path:
    """Fixture 2: headings like I., A., B., II., A."""
    doc = Document()
    doc.add_paragraph("Nama : Jane")
    doc.add_paragraph("NIM : 24100001")

    p = doc.add_paragraph("I. PENDAHULUAN")
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph("A. Tujuan")
    p.runs[0].bold = True
    doc.add_paragraph("Isi tujuan lama")
    p = doc.add_paragraph("B. Dasar Teori")
    p.runs[0].bold = True
    doc.add_paragraph("Isi dasar teori lama")

    p = doc.add_paragraph("II. HASIL")
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph("A. Pengujian")
    p.runs[0].bold = True
    doc.add_paragraph("Isi pengujian lama")
    doc.save(path)
    return path


def build_table_docx(path: Path) -> Path:
    """Fixture 5-ish: identity table + data table."""
    doc = Document()
    ident = doc.add_table(rows=2, cols=2)
    ident.cell(0, 0).text = "Nama"
    ident.cell(0, 1).text = "John Doe"
    ident.cell(1, 0).text = "NIM"
    ident.cell(1, 1).text = "24123456"

    doc.add_heading("BAB I", level=1)
    doc.add_heading("HASIL", level=1)
    tbl = doc.add_table(rows=3, cols=3)
    headers = ["No", "Pengujian", "Hasil"]
    for i, h in enumerate(headers):
        tbl.cell(0, i).text = h
    tbl.cell(1, 0).text = "1"
    tbl.cell(1, 1).text = "Login"
    tbl.cell(1, 2).text = "Berhasil"
    tbl.cell(2, 0).text = "2"
    tbl.cell(2, 1).text = "Logout"
    tbl.cell(2, 2).text = "Berhasil"
    doc.save(path)
    return path
