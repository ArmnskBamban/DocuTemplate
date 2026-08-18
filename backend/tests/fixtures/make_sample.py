"""CLI/dev helper: generate a realistic sample DOCX for manual testing.

Usage: uv run python tests/fixtures/make_sample.py [output_path]
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm


def main() -> None:
    out = Path(sys.argv[1]) if len(sys.argv) > 1 else Path("sample.docx")
    d = Document()
    d.add_paragraph("UNIVERSITAS ANDALAS").alignment = WD_ALIGN_PARAGRAPH.CENTER
    d.add_paragraph("LAPORAN PRAKTIKUM").alignment = WD_ALIGN_PARAGRAPH.CENTER
    d.add_paragraph("Nama : Test User")
    d.add_paragraph("NIM : 24999999")
    d.add_heading("BAB I", level=1)
    d.add_heading("PENDAHULUAN", level=1)
    d.add_heading("1.1 Latar Belakang", level=2)
    d.add_paragraph("Perkembangan teknologi informasi pada saat ini mengalami perkembangan yang sangat pesat.")
    d.add_heading("1.2 Tujuan", level=2)
    d.add_paragraph("1. Mengetahui cara kerja sistem.")
    d.add_heading("BAB II", level=1)
    d.add_heading("LANDASAN TEORI", level=1)
    d.add_paragraph("Routing adalah suatu proses untuk meneruskan paket data.")
    d.sections[0].top_margin = Cm(3)
    d.sections[0].left_margin = Cm(4)
    d.save(out)
    print(f"Sample DOCX created: {out}")


if __name__ == "__main__":
    main()
